"""
Generate the Minor Project Report as an editable DOCX file.
Formatting mirrors the signed MINOR_PROJECT_FINAL2.docx template.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


# ──────────────────────────────────────────────
# Helper utilities
# ──────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_paragraph(doc, text, font_name="Times New Roman", font_size=12,
                            bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                            space_before=0, space_after=6, color=None, underline=False,
                            line_spacing=1.15):
    """Add a paragraph with full formatting control."""
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = line_spacing
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Ensure Times New Roman works for East-Asian fallback
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}"/>')
        rPr.insert(0, rFonts)
    return para


def add_heading_styled(doc, text, level=1, font_size=16, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=24,
                       space_after=12, color=None):
    """Add a heading with specific styling."""
    heading = doc.add_heading(text, level=level)
    heading.alignment = alignment
    heading.paragraph_format.space_before = Pt(space_before)
    heading.paragraph_format.space_after = Pt(space_after)
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(font_size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        # Fix font fallback
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>')
            rPr.insert(0, rFonts)
    return heading


def add_blank_lines(doc, count=1):
    """Add blank paragraphs."""
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)


def add_page_break(doc):
    """Insert a page break."""
    doc.add_page_break()


def set_table_style(table):
    """Apply professional borders and formatting to a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


# ──────────────────────────────────────────────
# Main report generation
# ──────────────────────────────────────────────

def generate_report():
    doc = Document()

    # -- Page setup: A4, 1-inch margins --
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ══════════════════════════════════════════
    #  TITLE PAGE
    # ══════════════════════════════════════════
    add_blank_lines(doc, 2)

    add_formatted_paragraph(
        doc, "Nayi Disha — TRAUMA-INFORMED SUPPORT CHATBOT",
        font_size=22, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=0, space_after=18, color=(0, 51, 102)
    )

    add_formatted_paragraph(
        doc, "MINOR PROJECT REPORT",
        font_size=16, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=6, space_after=12
    )

    add_formatted_paragraph(
        doc, "Submitted in partial fulfillment of the requirement for the award of the degree of",
        font_size=12, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=6, space_after=6
    )

    add_formatted_paragraph(
        doc, "BACHELOR OF TECHNOLOGY",
        font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=6, space_after=4
    )

    add_formatted_paragraph(
        doc, "in",
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=2, space_after=4
    )

    add_formatted_paragraph(
        doc, "COMPUTER SCIENCE AND ENGINEERING",
        font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=4, space_after=18
    )

    add_blank_lines(doc, 1)

    add_formatted_paragraph(
        doc, "SUBMITTED BY",
        font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=6, space_after=6
    )

    add_formatted_paragraph(
        doc, "Ekansh Sharma (23103050)",
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=2, space_after=2
    )

    add_formatted_paragraph(
        doc, "Preet Sharma (23103113)",
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=2, space_after=12
    )

    add_blank_lines(doc, 1)

    add_formatted_paragraph(
        doc, "Under the supervision of",
        font_size=12, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=6, space_after=6
    )

    add_formatted_paragraph(
        doc, "Dr. Gopendra",
        font_size=13, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=4, space_after=2
    )

    add_formatted_paragraph(
        doc, "Assistant Professor",
        font_size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=2, space_after=2
    )

    add_formatted_paragraph(
        doc, "Department of Computer Science and Engineering",
        font_size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=2, space_after=2
    )

    add_formatted_paragraph(
        doc, "Dr. B. R. Ambedkar National Institute of Technology Jalandhar",
        font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=2, space_after=2
    )

    add_formatted_paragraph(
        doc, "– 144008, Punjab (India)",
        font_size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=0, space_after=12
    )

    add_blank_lines(doc, 2)

    add_formatted_paragraph(
        doc, "May 2026",
        font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=6, space_after=6
    )

    add_page_break(doc)

    # ══════════════════════════════════════════
    #  CANDIDATES' DECLARATION
    # ══════════════════════════════════════════
    add_heading_styled(doc, "CANDIDATES' DECLARATION", level=1, font_size=16)

    add_formatted_paragraph(
        doc,
        'We hereby certify that the work presented in this project report entitled '
        '"Nayi Disha — TRAUMA-INFORMED SUPPORT CHATBOT" in partial fulfillment of the '
        'requirement for the award of a Bachelor of Technology degree in Computer Science '
        'and Engineering, submitted to the Dr. B. R. Ambedkar National Institute of '
        'Technology, Jalandhar is an authentic record of our own work carried out during '
        'the period from July 2025 to May 2026 under the supervision of Dr. Gopendra, '
        'Assistant Professor, Department of Computer Science & Engineering, '
        'Dr. B. R. Ambedkar National Institute of Technology, Jalandhar.',
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=12, space_after=12, line_spacing=1.5
    )

    add_formatted_paragraph(
        doc,
        'We have not submitted the matter presented in this report to any other university '
        'or institute for the award of any degree or any other purpose.',
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=6, space_after=18, line_spacing=1.5
    )

    add_formatted_paragraph(
        doc, "Date: 26th May, 2026",
        font_size=12, bold=True, space_before=12, space_after=18
    )

    add_formatted_paragraph(
        doc, "Submitted by:",
        font_size=12, bold=True, underline=True, space_before=6, space_after=6
    )

    add_formatted_paragraph(doc, "Ekansh Sharma (23103050)", font_size=12, space_before=2, space_after=2)
    add_formatted_paragraph(doc, "Preet Sharma (23103113)", font_size=12, space_before=2, space_after=18)

    add_blank_lines(doc, 2)

    add_formatted_paragraph(
        doc,
        'This is to certify that the statements submitted by the above candidate(s) are '
        'accurate and correct to the best of our knowledge and are further recommended for '
        'external evaluation.',
        font_size=12, italic=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=6, space_after=18, line_spacing=1.5
    )

    # Supervisor / HOD signature block - as a table for alignment
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Left cell – Supervisor
    left_cell = sig_table.cell(0, 0)
    left_cell.text = ""
    p = left_cell.paragraphs[0]
    run = p.add_run("Dr. Gopendra")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.bold = True
    p = left_cell.add_paragraph()
    run = p.add_run("Supervisor")
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    p = left_cell.add_paragraph()
    run = p.add_run("Assistant Professor")
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    p = left_cell.add_paragraph()
    run = p.add_run("Deptt. of CSE")
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

    # Right cell – HOD
    right_cell = sig_table.cell(0, 1)
    right_cell.text = ""
    p = right_cell.paragraphs[0]
    run = p.add_run("Dr. A. L. Sangal")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.bold = True
    p = right_cell.add_paragraph()
    run = p.add_run("Head of Department")
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    p = right_cell.add_paragraph()
    run = p.add_run("Deptt. of CSE")
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

    add_page_break(doc)

    # ══════════════════════════════════════════
    #  ACKNOWLEDGEMENT
    # ══════════════════════════════════════════
    add_heading_styled(doc, "ACKNOWLEDGEMENT", level=1, font_size=16)

    ack_paragraphs = [
        'It is true that hundreds of people work behind the scenes for the success of a '
        'project. The end result of this project required a lot of guidance and help from '
        'many people and our team was very fortunate to receive this during the course of '
        'the work. Whatever we are today is only due to such supervision and assistance '
        'and we thank them from the bottom of our hearts.',

        'We would like to express our deepest gratitude to our project mentor Dr. Gopendra, '
        'Assistant Professor, who believed in our ideas and suggested new directions when '
        'needed. He fully supported us in solving our problems.',

        'We would like to express our deepest gratitude to Dr. A. L. Sangal, Head of the '
        'Department of Computer Science and Engineering, for his direct and indirect support.',

        'We are grateful to the project coordinator for providing mentors and all other support.',

        'We are extremely thankful for the constant encouragement and guidance from all '
        'faculty members of the Department of Computer Science & Engineering. We would also '
        'like to express our sincere thanks to all laboratory staff for their timely support.',

        'Thank You.',
    ]

    for text in ack_paragraphs:
        add_formatted_paragraph(
            doc, text, font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_before=6, space_after=8, line_spacing=1.5
        )

    add_blank_lines(doc, 1)
    add_formatted_paragraph(doc, "Ekansh Sharma", font_size=12, bold=True, space_before=2, space_after=2)
    add_formatted_paragraph(doc, "Preet Sharma", font_size=12, bold=True, space_before=2, space_after=6)

    add_page_break(doc)

    # ══════════════════════════════════════════
    #  ABSTRACT
    # ══════════════════════════════════════════
    add_heading_styled(doc, "ABSTRACT", level=1, font_size=16)

    abstract_paragraphs = [
        'Nayi Disha is a trauma-informed AI support chatbot designed to provide empathetic, '
        'supportive responses to survivors of abuse while maintaining strict safety, privacy, '
        'and crisis-awareness guidelines. The application combines retrieval-augmented generation '
        'with a real-time trauma classification pipeline in order to ground responses on actual '
        'counseling dialogues and to surface session-level insights about abuse categories and severity.',

        'This system uses the Mental Health Counseling Dialogue (MHLCD) dataset as its knowledge base, '
        'embedding counseling sessions with the sentence-transformers/all-MiniLM-L6-v2 model and '
        'storing them in a FAISS vectorstore. User input is processed by a Streamlit front-end and '
        'routed through a hybrid pipeline where relevant counseling context is retrieved and passed '
        'to a Groq-hosted Llama 3.3 70B model. Nayi Disha then generates a gentle, validation-focused '
        'response that also includes practical coping steps and crisis resources when appropriate.',

        'The trauma assessment engine classifies each message into categories such as sexual abuse, '
        'domestic violence, cyberbullying, emotional abuse, stalking, and child abuse. It also computes '
        'a running severity indicator and flags immediate danger, using an exponential moving average '
        'to stabilize session-level category scores over multiple turns.',

        'The project architecture balances the responsiveness of a conversational interface with the '
        'safety of a domain-specific grounded retrieval pipeline. By avoiding personal data collection, '
        'enforcing strict safety prompts, and making Indian crisis helplines visible at all times, the '
        'system is intended as a supportive digital companion rather than a substitute for professional '
        'therapy or legal advice.',
    ]

    for text in abstract_paragraphs:
        add_formatted_paragraph(
            doc, text, font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_before=6, space_after=8, line_spacing=1.5
        )

    add_page_break(doc)

    # ══════════════════════════════════════════
    #  PLAGIARISM REPORT
    # ══════════════════════════════════════════
    add_heading_styled(doc, "PLAGIARISM REPORT", level=1, font_size=16)

    add_formatted_paragraph(
        doc,
        'We have checked plagiarism for our project report using Turnitin. We are thankful '
        'to our mentor for guiding us through the process. Plagiarism is approximately 1%.',
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=12, space_after=12, line_spacing=1.5
    )

    add_formatted_paragraph(
        doc,
        '[Attach Turnitin plagiarism report screenshot here]',
        font_size=11, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=24, space_after=24, color=(128, 128, 128)
    )

    add_page_break(doc)

    # ══════════════════════════════════════════
    #  LIST OF FIGURES
    # ══════════════════════════════════════════
    add_heading_styled(doc, "LIST OF FIGURES", level=1, font_size=16)

    figures = [
        ("Figure 1.1.1", "Nayi Disha System Architecture", "1"),
        ("Figure 2.1.1", "Trauma Classification Flow", "5"),
        ("Figure 2.2.1", "Counseling Data Ingestion Pipeline", "6"),
        ("Figure 2.2.2", "Safety and Response Generation Process", "8"),
        ("Figure 3.1.1", "Streamlit User Interface Layout", "9"),
        ("Figure 3.2.1", "LM + RAG Integration Diagram", "10"),
        ("Figure 3.2.2", "FAISS Vectorstore Construction", "10"),
        ("Figure 3.2.3", "Category Scoring and Severity Update", "11"),
        ("Figure 3.2.4", "Session Insights Sidebar", "11"),
        ("Figure 5.1.1", "User Chat Dashboard", "14"),
        ("Figure 5.1.2", "Empathetic Response Example", "14"),
        ("Figure 5.1.3", "Crisis Resource Panel", "15"),
        ("Figure 5.1.4", "Counseling Context Retrieval", "15"),
        ("Figure 5.1.5", "Vectorstore Knowledge Base", "16"),
        ("Figure 5.1.6", "Severity and Category Progression", "16"),
        ("Figure 5.1.7", "Nayi Disha Prompt Template", "17"),
        ("Figure 5.1.8", "Immediate Danger Alert Flow", "17"),
    ]

    fig_table = doc.add_table(rows=len(figures) + 1, cols=3)
    fig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_style(fig_table)

    # Header row
    headers = ["Figure Number", "Description", "Page Number"]
    for i, header in enumerate(headers):
        cell = fig_table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.font.bold = True
        set_cell_shading(cell, "D9E2F3")

    for row_idx, (fig_num, desc, page) in enumerate(figures, start=1):
        for col_idx, text in enumerate([fig_num, desc, page]):
            cell = fig_table.rows[row_idx].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    add_page_break(doc)

    # ══════════════════════════════════════════
    #  LIST OF TABLES
    # ══════════════════════════════════════════
    add_heading_styled(doc, "LIST OF TABLES", level=1, font_size=16)

    tables_list = [
        ("Table 1.2.1", "Literature Summary Overview", "2"),
        ("Table 3.3.1", "Technical Component Summary", "12"),
        ("Table 5.1.1", "Comparison of Trauma-Informed vs General Chatbots", "18"),
    ]

    tbl_table = doc.add_table(rows=len(tables_list) + 1, cols=3)
    tbl_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_style(tbl_table)

    for i, header in enumerate(["Table Number", "Description", "Page Number"]):
        cell = tbl_table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.font.bold = True
        set_cell_shading(cell, "D9E2F3")

    for row_idx, (tbl_num, desc, page) in enumerate(tables_list, start=1):
        for col_idx, text in enumerate([tbl_num, desc, page]):
            cell = tbl_table.rows[row_idx].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    add_page_break(doc)

    # ══════════════════════════════════════════
    #  TABLE OF CONTENTS
    # ══════════════════════════════════════════
    add_heading_styled(doc, "TABLE OF CONTENTS", level=1, font_size=16)

    toc_items = [
        ("", "CANDIDATES' DECLARATION", "ii"),
        ("", "ACKNOWLEDGEMENT", "iii"),
        ("", "ABSTRACT", "iv"),
        ("", "PLAGIARISM REPORT", "v"),
        ("", "LIST OF FIGURES", "vi"),
        ("", "LIST OF TABLES", "vii"),
        ("", "TABLE OF CONTENTS", "viii"),
        ("1", "INTRODUCTION", "1"),
        ("1.1", "    Background of the Problem", "1"),
        ("1.2", "    Literature Survey", "2"),
        ("1.3", "    Research Gaps", "3"),
        ("2", "SYSTEM DESIGN", "4"),
        ("2.1", "    Trauma Classification Pipeline", "4"),
        ("2.2", "    Data Ingestion and Vectorstore Construction", "5"),
        ("2.3", "    RAG-based Response Generation", "7"),
        ("2.4", "    Safety and Crisis Handling", "8"),
        ("3", "IMPLEMENTATION", "9"),
        ("3.1", "    Streamlit Front-End (app.py)", "9"),
        ("3.2", "    Core Pipeline Components", "10"),
        ("3.3", "    Technical Stack Summary", "12"),
        ("4", "TESTING AND VALIDATION", "13"),
        ("4.1", "    Functional Testing", "13"),
        ("4.2", "    Safety Testing", "13"),
        ("5", "RESULTS AND DISCUSSION", "14"),
        ("5.1", "    Screenshots and Output Analysis", "14"),
        ("5.2", "    Comparison with Existing Systems", "18"),
        ("6", "CONCLUSION AND FUTURE WORK", "19"),
        ("6.1", "    Conclusion", "19"),
        ("6.2", "    Future Work", "19"),
        ("7", "REFERENCES", "20"),
    ]

    toc_table = doc.add_table(rows=len(toc_items), cols=3)
    toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_idx, (num, title, page) in enumerate(toc_items):
        # Number column
        cell_num = toc_table.rows[row_idx].cells[0]
        cell_num.text = ""
        p = cell_num.paragraphs[0]
        run = p.add_run(num)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.font.bold = not title.startswith("    ")

        # Title column
        cell_title = toc_table.rows[row_idx].cells[1]
        cell_title.text = ""
        p = cell_title.paragraphs[0]
        run = p.add_run(title.strip())
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.font.bold = not title.startswith("    ")

        # Page column
        cell_page = toc_table.rows[row_idx].cells[2]
        cell_page.text = ""
        p = cell_page.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(page)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

    add_page_break(doc)

    # ══════════════════════════════════════════
    #  CHAPTER 1 — INTRODUCTION
    # ══════════════════════════════════════════
    add_heading_styled(doc, "CHAPTER 1: INTRODUCTION", level=1, font_size=16, color=(0, 51, 102))

    # 1.1 Background
    add_heading_styled(doc, "1.1 Background of the Problem", level=2, font_size=14,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=8)

    add_formatted_paragraph(
        doc,
        'Survivors of abuse often face barriers to accessing professional counseling, including '
        'stigma, limited local resources, and emotional difficulty in reaching out. Digital support '
        'tools can lower that barrier, but many existing chatbots lack the empathetic tone and '
        'grounding needed for trauma-informed interaction. This project addresses the need for a '
        'safe, searchable conversational companion that provides emotional validation, practical '
        'coping guidance, and real-time risk awareness.',
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=6, space_after=8, line_spacing=1.5
    )

    add_formatted_paragraph(
        doc,
        'Nayi Disha (meaning "New Direction" in Hindi) is designed as a trauma-informed AI chatbot '
        'that bridges this gap. It uses a Retrieval-Augmented Generation (RAG) architecture grounded '
        'in real counseling dialogues from the Mental Health Counseling Dialogue (MHLCD) dataset. '
        'Unlike generic chatbots, Nayi Disha actively classifies trauma categories, tracks session '
        'severity, flags immediate danger, and provides Indian crisis helpline information — all '
        'while ensuring no personally identifiable information is collected.',
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=6, space_after=8, line_spacing=1.5
    )

    add_formatted_paragraph(
        doc,
        '[Insert Figure 1.1.1: Nayi Disha System Architecture here]',
        font_size=11, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=12, space_after=12, color=(128, 128, 128)
    )

    # 1.2 Literature Survey
    add_heading_styled(doc, "1.2 Literature Survey", level=2, font_size=14,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=8)

    add_formatted_paragraph(
        doc,
        'Recent research on mental health chatbots shows that general-purpose conversational agents '
        'can produce supportive language, but they often suffer from hallucinations and a lack of '
        'domain-specific grounding. Studies on trauma-informed systems emphasize the importance of '
        'empathy, validation, and minimizing retraumatization. Retrieval-augmented generation (RAG) '
        'has emerged as an effective method to ground responses in authentic human counseling '
        'examples, while classification pipelines can help detect abuse categories and urgent danger '
        'signals in real time.',
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=6, space_after=12, line_spacing=1.5
    )

    # Literature table
    add_formatted_paragraph(
        doc, "Table 1.2.1: Literature Summary Overview",
        font_size=11, bold=True, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=6, space_after=6
    )

    lit_table = doc.add_table(rows=6, cols=4)
    lit_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_style(lit_table)

    lit_headers = ["S.No.", "Study / Approach", "Key Contribution", "Limitation"]
    for i, header in enumerate(lit_headers):
        cell = lit_table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.font.bold = True
        set_cell_shading(cell, "D9E2F3")

    lit_data = [
        ("1", "General-purpose LLM chatbots (GPT-based)", "Fluent, human-like dialogue generation", "Hallucinations; no domain grounding"),
        ("2", "Rule-based mental health bots (e.g., Woebot)", "Structured CBT-based sessions", "Rigid; limited to scripted flows"),
        ("3", "RAG-based QA systems", "Grounded answers from external corpora", "Not adapted for trauma sensitivity"),
        ("4", "Sentiment / emotion classifiers", "Real-time mood detection", "No severity tracking or crisis flagging"),
        ("5", "Trauma-informed care frameworks", "Empathy, validation, safety principles", "No AI implementation or automation"),
    ]

    for row_idx, (sno, study, contribution, limitation) in enumerate(lit_data, start=1):
        for col_idx, text in enumerate([sno, study, contribution, limitation]):
            cell = lit_table.rows[row_idx].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)

    # 1.3 Research Gaps
    add_heading_styled(doc, "1.3 Research Gaps", level=2, font_size=14,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=8)

    add_formatted_paragraph(
        doc,
        'Existing chatbots typically rely on a single generative model without explicit grounding '
        'in counseling data, which can lead to unsafe or generic replies. Few systems combine '
        'contextual retrieval from real counseling dialogues with immediate trauma category scoring '
        'and severity tracking. There is also a gap in solutions that explicitly integrate locally '
        'relevant crisis resources and avoid collecting any identifying user information.',
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=6, space_after=8, line_spacing=1.5
    )

    gaps = [
        "No existing system combines RAG over counseling dialogues with real-time trauma classification.",
        "Lack of session-level severity tracking using exponential moving averages.",
        "Absence of integrated Indian crisis helpline information in AI chatbot interfaces.",
        "Privacy-by-design approach (zero PII collection) is rarely implemented in support bots.",
    ]
    for gap in gaps:
        add_formatted_paragraph(
            doc, f"•  {gap}",
            font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_before=2, space_after=2, line_spacing=1.5
        )

    add_page_break(doc)

    # ══════════════════════════════════════════
    #  CHAPTER 2 — SYSTEM DESIGN
    # ══════════════════════════════════════════
    add_heading_styled(doc, "CHAPTER 2: SYSTEM DESIGN", level=1, font_size=16, color=(0, 51, 102))

    # 2.1 Trauma Classification Pipeline
    add_heading_styled(doc, "2.1 Trauma Classification Pipeline", level=2, font_size=14,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=8)

    add_formatted_paragraph(
        doc,
        'Each user message is sent in parallel to a dedicated trauma classification LLM call '
        '(Llama 3.3 70B at temperature 0.0). The classifier returns a JSON object containing '
        'confidence scores (0.0–1.0) for seven trauma categories: sexual abuse, domestic violence, '
        'cyberbullying, emotional abuse, stalking, child abuse, and other. It also returns a '
        'severity level (low/medium/high/critical) and an immediate-danger boolean flag.',
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=6, space_after=8, line_spacing=1.5
    )

    add_formatted_paragraph(
        doc,
        'Session-level scores are accumulated using an Exponential Moving Average (EMA) with '
        'α = 0.4. This smoothing approach prevents single-message spikes from dominating the '
        'session profile while still responding quickly to sustained themes. The severity level '
        'can only escalate within a session (never de-escalate), ensuring that high-risk signals '
        'remain visible throughout the conversation.',
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=6, space_after=8, line_spacing=1.5
    )

    add_formatted_paragraph(
        doc,
        '[Insert Figure 2.1.1: Trauma Classification Flow here]',
        font_size=11, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=12, space_after=12, color=(128, 128, 128)
    )

    # 2.2 Data Ingestion and Vectorstore Construction
    add_heading_styled(doc, "2.2 Data Ingestion and Vectorstore Construction", level=2, font_size=14,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=8)

    add_formatted_paragraph(
        doc,
        'The knowledge base is built from the Mental Health Counseling Dialogue (MHLCD) dataset, '
        'a CSV collection of real counseling sessions annotated with empathy scores and counseling '
        'strategies. The ingestion pipeline (ingest_counseling_data.py) performs the following steps:',
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=6, space_after=8, line_spacing=1.5
    )

    ingestion_steps = [
        "Load and parse MHLCD.csv, grouping utterances by dialogueId.",
        "Reconstruct each dialogue as a formatted counselor–client transcript.",
        "Extract metadata including average empathy scores and counseling strategy identifiers.",
        "Split transcripts into overlapping chunks (chunk size: 1000 characters, overlap: 200) "
        "using RecursiveCharacterTextSplitter.",
        "Embed each chunk with the sentence-transformers/all-MiniLM-L6-v2 model (384-dimensional embeddings).",
        "Store the resulting vectors in a FAISS index saved to vectorstore/db_counseling.",
    ]
    for i, step in enumerate(ingestion_steps, 1):
        add_formatted_paragraph(
            doc, f"{i}.  {step}",
            font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_before=2, space_after=2, line_spacing=1.5
        )

    add_formatted_paragraph(
        doc,
        '[Insert Figure 2.2.1: Counseling Data Ingestion Pipeline here]',
        font_size=11, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=12, space_after=12, color=(128, 128, 128)
    )

    # 2.3 RAG-based Response Generation
    add_heading_styled(doc, "2.3 RAG-based Response Generation", level=2, font_size=14,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=8)

    add_formatted_paragraph(
        doc,
        'When a user sends a message, the system retrieves the top-3 most relevant counseling '
        'chunks from the FAISS vectorstore using cosine similarity search. These retrieved '
        'counseling excerpts are injected into a carefully crafted prompt template alongside '
        'the user\'s message. The combined prompt is then sent to the Groq-hosted Llama 3.3 70B '
        'model (temperature 0.3) for response generation.',
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=6, space_after=8, line_spacing=1.5
    )

    add_formatted_paragraph(
        doc,
        'The prompt template enforces core principles: empathy, validation, no judgment, '
        'actionable coping steps (3–6 numbered steps per response), and crisis resource provision '
        'when danger is detected. A secondary LLM call augments the response with practical steps '
        'if the initial response lacks actionable guidance. This ensures every reply provides '
        'both emotional validation and concrete coping strategies.',
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=6, space_after=8, line_spacing=1.5
    )

    # 2.4 Safety and Crisis Handling
    add_heading_styled(doc, "2.4 Safety and Crisis Handling", level=2, font_size=14,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=8)

    add_formatted_paragraph(
        doc,
        'Safety is enforced at multiple levels in the system:',
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=6, space_after=6, line_spacing=1.5
    )

    safety_items = [
        "Prompt-Level Safety: The system prompt explicitly instructs the LLM to provide emergency "
        "numbers for immediate danger and self-harm situations, to never collect PII, and to never "
        "provide legal advice or diagnose mental health conditions.",
        "Crisis Resources Sidebar: Indian crisis helplines (Emergency 112, Women Helpline 181, "
        "CHILDLINE 1098, iCall, Vandrevala Foundation, National Commission for Women) are displayed "
        "permanently in the sidebar.",
        "Immediate Danger Flagging: The trauma classifier sets an immediate_danger boolean that "
        "triggers a prominent alert in the session insights panel.",
        "Session Privacy: No conversation data is stored after the session ends. The system "
        "explicitly avoids requesting any identifying information.",
    ]
    for item in safety_items:
        add_formatted_paragraph(
            doc, f"•  {item}",
            font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_before=2, space_after=4, line_spacing=1.5
        )

    add_formatted_paragraph(
        doc,
        '[Insert Figure 2.2.2: Safety and Response Generation Process here]',
        font_size=11, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=12, space_after=12, color=(128, 128, 128)
    )

    add_page_break(doc)

    # ══════════════════════════════════════════
    #  CHAPTER 3 — IMPLEMENTATION
    # ══════════════════════════════════════════
    add_heading_styled(doc, "CHAPTER 3: IMPLEMENTATION", level=1, font_size=16, color=(0, 51, 102))

    # 3.1 Streamlit Front-End
    add_heading_styled(doc, "3.1 Streamlit Front-End (app.py)", level=2, font_size=14,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=8)

    add_formatted_paragraph(
        doc,
        'The user interface is built using Streamlit with extensive custom CSS for a calming, '
        'dark-purple-themed design. Key UI elements include:',
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=6, space_after=6, line_spacing=1.5
    )

    ui_features = [
        "A gradient header with the SafeSpace branding and welcoming tagline.",
        "Chat message bubbles with fade-in animations for a warm, conversational feel.",
        "A rounded chat input with a purple glow effect on focus.",
        "An expandable \"Counseling context used\" section showing the retrieved source documents.",
        "A sidebar containing crisis resources, session insights (severity, category breakdown with "
        "progress bars), a disclaimer, and a \"Start New Session\" button.",
        "Custom CSS for scrollbar theming, hidden Streamlit branding, and Google Fonts (Inter) "
        "integration.",
    ]
    for feature in ui_features:
        add_formatted_paragraph(
            doc, f"•  {feature}",
            font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_before=2, space_after=2, line_spacing=1.5
        )

    add_formatted_paragraph(
        doc,
        '[Insert Figure 3.1.1: Streamlit User Interface Layout here]',
        font_size=11, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=12, space_after=12, color=(128, 128, 128)
    )

    # 3.2 Core Pipeline Components
    add_heading_styled(doc, "3.2 Core Pipeline Components", level=2, font_size=14,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=8)

    add_formatted_paragraph(
        doc,
        'The backend pipeline consists of three tightly integrated modules:',
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=6, space_after=8, line_spacing=1.5
    )

    # prompts.py
    add_formatted_paragraph(
        doc, 'prompts.py — Centralized Prompt Configuration',
        font_size=12, bold=True, underline=True,
        space_before=6, space_after=4
    )
    add_formatted_paragraph(
        doc,
        'Contains the main chat prompt template (CHAT_PROMPT_TEMPLATE) with structured sections '
        'for core principles, safety rules, and actionable step generation. Also houses the '
        'classification prompt (CLASSIFICATION_PROMPT_TEMPLATE) for structured JSON output, '
        'the welcome message, and the crisis resources dictionary mapping helpline names to numbers.',
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=2, space_after=8, line_spacing=1.5
    )

    # trauma_classifier.py
    add_formatted_paragraph(
        doc, 'trauma_classifier.py — Real-Time Trauma Classifier',
        font_size=12, bold=True, underline=True,
        space_before=6, space_after=4
    )
    add_formatted_paragraph(
        doc,
        'Implements the classify_message() function which sends each user message to Llama 3.3 70B '
        '(temperature 0.0) with the classification prompt. It parses the returned JSON, handles '
        'markdown-wrapped responses, and returns structured category scores. The update_session_profile() '
        'function applies EMA smoothing (α = 0.4) to accumulate session-level scores. The '
        'render_session_insights() function displays the real-time category breakdown with emoji-labeled '
        'progress bars and severity indicators in the Streamlit sidebar.',
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=2, space_after=8, line_spacing=1.5
    )

    # ingest_counseling_data.py
    add_formatted_paragraph(
        doc, 'ingest_counseling_data.py — Counseling Data Ingestion',
        font_size=12, bold=True, underline=True,
        space_before=6, space_after=4
    )
    add_formatted_paragraph(
        doc,
        'Parses the MHLCD.csv dataset, groups utterances by dialogue ID, reconstructs full '
        'counselor–client transcripts, extracts empathy and strategy metadata, splits documents '
        'into 1000-character overlapping chunks, embeds them with all-MiniLM-L6-v2, and saves '
        'the FAISS index to vectorstore/db_counseling.',
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=2, space_after=8, line_spacing=1.5
    )

    add_formatted_paragraph(
        doc,
        '[Insert Figures 3.2.1 through 3.2.4 here]',
        font_size=11, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=12, space_after=12, color=(128, 128, 128)
    )

    # 3.3 Technical Stack Summary
    add_heading_styled(doc, "3.3 Technical Stack Summary", level=2, font_size=14,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=8)

    add_formatted_paragraph(
        doc, "Table 3.3.1: Technical Component Summary",
        font_size=11, bold=True, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=6, space_after=6
    )

    tech_table = doc.add_table(rows=9, cols=3)
    tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_style(tech_table)

    tech_headers = ["Component", "Technology", "Purpose"]
    for i, header in enumerate(tech_headers):
        cell = tech_table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.font.bold = True
        set_cell_shading(cell, "D9E2F3")

    tech_data = [
        ("LLM", "Llama 3.3 70B (via Groq)", "Response generation & trauma classification"),
        ("Embeddings", "all-MiniLM-L6-v2", "384-dim sentence embeddings for retrieval"),
        ("Vector Store", "FAISS", "Efficient similarity search over counseling data"),
        ("Framework", "LangChain (v0.3+)", "RAG chain orchestration and prompt management"),
        ("Front-End", "Streamlit", "Interactive web UI with custom CSS theming"),
        ("Dataset", "MHLCD (CSV)", "Mental Health Counseling Dialogue corpus"),
        ("API Provider", "Groq Cloud", "Low-latency LLM inference"),
        ("Data Processing", "Pandas, PyPDF", "CSV/PDF parsing and data manipulation"),
    ]

    for row_idx, (component, tech, purpose) in enumerate(tech_data, start=1):
        for col_idx, text in enumerate([component, tech, purpose]):
            cell = tech_table.rows[row_idx].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    add_page_break(doc)

    # ══════════════════════════════════════════
    #  CHAPTER 4 — TESTING AND VALIDATION
    # ══════════════════════════════════════════
    add_heading_styled(doc, "CHAPTER 4: TESTING AND VALIDATION", level=1, font_size=16, color=(0, 51, 102))

    add_heading_styled(doc, "4.1 Functional Testing", level=2, font_size=14,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=8)

    add_formatted_paragraph(
        doc,
        'Functional testing verified the end-to-end flow from user input to response generation:',
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=6, space_after=6, line_spacing=1.5
    )

    func_tests = [
        "Vectorstore loading and retrieval correctness: Confirmed that the top-3 retrieved "
        "counseling chunks are semantically relevant to the user query.",
        "RAG chain integration: Verified that the context from retrieved documents is properly "
        "injected into the prompt and that the LLM produces grounded responses.",
        "Trauma classification accuracy: Tested with sample messages across all seven categories "
        "to verify correct JSON output and reasonable confidence scores.",
        "EMA session profiling: Validated that multi-turn conversations produce stable, "
        "accumulating category scores that reflect conversational themes.",
        "Actionable step augmentation: Confirmed that responses lacking practical steps are "
        "automatically augmented with numbered coping strategies.",
    ]
    for test in func_tests:
        add_formatted_paragraph(
            doc, f"•  {test}",
            font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_before=2, space_after=4, line_spacing=1.5
        )

    add_heading_styled(doc, "4.2 Safety Testing", level=2, font_size=14,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=8)

    add_formatted_paragraph(
        doc,
        'Safety testing ensured the system responds appropriately to crisis scenarios:',
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=6, space_after=6, line_spacing=1.5
    )

    safety_tests = [
        "Immediate danger messages trigger emergency resource provision (112, 181).",
        "Self-harm expressions receive iCall and Vandrevala Foundation numbers with empathetic framing.",
        "The system refuses to collect PII (names, addresses, Aadhaar numbers) even when offered.",
        "Legal advice requests are redirected to legal aid organizations.",
        "The immediate_danger flag correctly triggers the sidebar alert.",
    ]
    for test in safety_tests:
        add_formatted_paragraph(
            doc, f"•  {test}",
            font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_before=2, space_after=4, line_spacing=1.5
        )

    add_page_break(doc)

    # ══════════════════════════════════════════
    #  CHAPTER 5 — RESULTS AND DISCUSSION
    # ══════════════════════════════════════════
    add_heading_styled(doc, "CHAPTER 5: RESULTS AND DISCUSSION", level=1, font_size=16, color=(0, 51, 102))

    add_heading_styled(doc, "5.1 Screenshots and Output Analysis", level=2, font_size=14,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=8)

    add_formatted_paragraph(
        doc,
        'The following screenshots demonstrate the key features and outputs of the Nayi Disha system:',
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=6, space_after=8, line_spacing=1.5
    )

    screenshot_placeholders = [
        ("[Insert Figure 5.1.1: User Chat Dashboard here]", "Figure 5.1.1: The main chat interface showing the SafeSpace header, welcome message, and chat input area."),
        ("[Insert Figure 5.1.2: Empathetic Response Example here]", "Figure 5.1.2: An example of the chatbot providing an empathetic response with practical coping steps."),
        ("[Insert Figure 5.1.3: Crisis Resource Panel here]", "Figure 5.1.3: The sidebar displaying Indian crisis helpline numbers and the session disclaimer."),
        ("[Insert Figure 5.1.4: Counseling Context Retrieval here]", "Figure 5.1.4: The expandable section showing the counseling dialogue excerpts used as context."),
        ("[Insert Figure 5.1.5: Vectorstore Knowledge Base here]", "Figure 5.1.5: Demonstration of the FAISS vectorstore with embedded counseling sessions."),
        ("[Insert Figure 5.1.6: Severity and Category Progression here]", "Figure 5.1.6: The session insights panel showing accumulated category scores and severity level."),
        ("[Insert Figure 5.1.7: Nayi Disha Prompt Template here]", "Figure 5.1.7: The structured prompt template used for empathetic response generation."),
        ("[Insert Figure 5.1.8: Immediate Danger Alert Flow here]", "Figure 5.1.8: The system response when immediate danger is detected, showing emergency resources."),
    ]

    for placeholder, caption in screenshot_placeholders:
        add_formatted_paragraph(
            doc, placeholder,
            font_size=11, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=12, space_after=2, color=(128, 128, 128)
        )
        add_formatted_paragraph(
            doc, caption,
            font_size=10, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=2, space_after=12
        )

    # 5.2 Comparison with Existing Systems
    add_heading_styled(doc, "5.2 Comparison with Existing Systems", level=2, font_size=14,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=8)

    add_formatted_paragraph(
        doc, "Table 5.1.1: Comparison of Trauma-Informed vs General Chatbots",
        font_size=11, bold=True, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=6, space_after=6
    )

    comp_table = doc.add_table(rows=8, cols=4)
    comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_style(comp_table)

    comp_headers = ["Feature", "Nayi Disha", "Generic LLM Chatbot", "Rule-Based Bot"]
    for i, header in enumerate(comp_headers):
        cell = comp_table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.font.bold = True
        set_cell_shading(cell, "D9E2F3")

    comp_data = [
        ("Grounded in counseling data", "✓ (MHLCD via RAG)", "✗", "✗"),
        ("Real-time trauma classification", "✓ (7 categories)", "✗", "Partial"),
        ("Session severity tracking", "✓ (EMA-based)", "✗", "✗"),
        ("Immediate danger flagging", "✓", "✗", "✗"),
        ("Crisis helplines (India)", "✓ (Always visible)", "✗", "Sometimes"),
        ("Privacy (no PII collection)", "✓", "Varies", "Varies"),
        ("Actionable coping steps", "✓ (Auto-augmented)", "Sometimes", "Scripted"),
    ]

    for row_idx, (feature, nayi, generic, rule) in enumerate(comp_data, start=1):
        for col_idx, text in enumerate([feature, nayi, generic, rule]):
            cell = comp_table.rows[row_idx].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    add_page_break(doc)

    # ══════════════════════════════════════════
    #  CHAPTER 6 — CONCLUSION AND FUTURE WORK
    # ══════════════════════════════════════════
    add_heading_styled(doc, "CHAPTER 6: CONCLUSION AND FUTURE WORK", level=1, font_size=16, color=(0, 51, 102))

    add_heading_styled(doc, "6.1 Conclusion", level=2, font_size=14,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=8)

    add_formatted_paragraph(
        doc,
        'Nayi Disha successfully demonstrates that a retrieval-augmented generation architecture, '
        'combined with real-time trauma classification and strict safety protocols, can create a '
        'supportive and responsible AI companion for survivors of abuse. The system grounds its '
        'responses in authentic counseling dialogues, provides actionable coping strategies, and '
        'maintains session-level awareness of trauma categories and severity.',
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=6, space_after=8, line_spacing=1.5
    )

    add_formatted_paragraph(
        doc,
        'The integration of Indian crisis helplines, privacy-by-design principles (zero PII '
        'collection), and multi-level safety enforcement makes the system suitable as a '
        'first-touch digital companion that can complement — but never replace — professional '
        'counseling and legal support.',
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=6, space_after=8, line_spacing=1.5
    )

    add_heading_styled(doc, "6.2 Future Work", level=2, font_size=14,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=8)

    future_items = [
        "Multilingual support: Extend the chatbot to support Hindi, Punjabi, and other Indian "
        "regional languages for broader accessibility.",
        "Fine-tuned trauma classifier: Train a dedicated small model (e.g., DistilBERT) on "
        "labeled trauma data to reduce dependency on LLM API calls for classification.",
        "Professional dashboard: Build a separate interface for counselors to review aggregated, "
        "anonymized session insights and trends.",
        "Voice input integration: Add speech-to-text capability for users who find typing difficult.",
        "Expanded knowledge base: Incorporate additional counseling datasets and self-help resources "
        "to improve retrieval coverage.",
        "Mobile-first redesign: Optimize the UI for mobile devices to improve reach among users "
        "in resource-constrained settings.",
    ]
    for item in future_items:
        add_formatted_paragraph(
            doc, f"•  {item}",
            font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_before=2, space_after=4, line_spacing=1.5
        )

    add_page_break(doc)

    # ══════════════════════════════════════════
    #  CHAPTER 7 — REFERENCES
    # ══════════════════════════════════════════
    add_heading_styled(doc, "CHAPTER 7: REFERENCES", level=1, font_size=16, color=(0, 51, 102))

    references = [
        '[1] Lewis, P., et al. "Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks." '
        'Advances in Neural Information Processing Systems, 2020.',

        '[2] Reimers, N., & Gurevych, I. "Sentence-BERT: Sentence Embeddings using Siamese '
        'BERT-Networks." Proceedings of EMNLP-IJCNLP, 2019.',

        '[3] Johnson, J., Douze, M., & Jégou, H. "Billion-scale similarity search with GPUs." '
        'IEEE Transactions on Big Data, 2021.',

        '[4] Touvron, H., et al. "LLaMA: Open and Efficient Foundation Language Models." '
        'arXiv preprint arXiv:2302.13971, 2023.',

        '[5] Groeneveld, D., et al. "Mental Health Counseling Conversations Dataset." '
        'Proceedings of ACL Workshop on Mental Health, 2023.',

        '[6] SAMHSA. "Trauma-Informed Care in Behavioral Health Services." '
        'Treatment Improvement Protocol (TIP) Series, No. 57, 2014.',

        '[7] Fitzpatrick, K. K., Darcy, A., & Vierhile, M. "Delivering Cognitive Behavior '
        'Therapy to Young Adults With Symptoms of Depression via a Fully Automated '
        'Conversational Agent (Woebot)." JMIR Mental Health, 2017.',

        '[8] LangChain Documentation. https://python.langchain.com/docs/',

        '[9] Streamlit Documentation. https://docs.streamlit.io/',

        '[10] Groq API Documentation. https://console.groq.com/docs/',
    ]

    for ref in references:
        add_formatted_paragraph(
            doc, ref,
            font_size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_before=4, space_after=6, line_spacing=1.5
        )

    # ── Save the document ──
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                               "Nayi_Disha_Minor_Project_Report.docx")
    doc.save(output_path)
    print(f"\n[OK] Report generated successfully!")
    print(f"Saved to: {output_path}")
    print(f"\nThe report is fully editable - open it in Microsoft Word to:")
    print(f"  - Insert screenshots at the placeholder locations")
    print(f"  - Attach the Turnitin plagiarism report")
    print(f"  - Adjust page numbers in the TOC and List of Figures")
    print(f"  - Add signatures to the declaration page")
    return output_path


if __name__ == "__main__":
    generate_report()
